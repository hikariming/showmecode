import pytest
from docx import Document
from docx.shared import Pt


@pytest.fixture
def make_doc():
    """Build an in-memory Document. `paras` is a list of (text, font_name | None)."""
    def _make(paras):
        doc = Document()
        for text, font in paras:
            p = doc.add_paragraph()
            run = p.add_run(text)
            if font:
                run.font.name = font
                run.font.size = Pt(11)
        return doc
    return _make
