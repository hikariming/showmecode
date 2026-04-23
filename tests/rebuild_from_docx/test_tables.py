from docx import Document

from scripts.rebuild_from_docx.tables import (
    classify_table,
    render_code_table,
    render_data_table,
    render_callout_table,
)


def _add_consolas_table(doc: Document, lang_first: bool, body: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    p = t.rows[0].cells[0].paragraphs[0]
    if lang_first:
        text = body
    else:
        text = body
    run = p.add_run(text)
    run.font.name = "Consolas"


def test_classify_code_container():
    doc = Document()
    _add_consolas_table(doc, lang_first=True, body="JSON\n{\"k\": 1}")
    assert classify_table(doc.tables[0]) == "code"


def test_classify_data_table():
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "h1"
    t.rows[0].cells[1].text = "h2"
    t.rows[1].cells[0].text = "v1"
    t.rows[1].cells[1].text = "v2"
    assert classify_table(doc.tables[0]) == "data"


def test_classify_callout():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "just a note"
    assert classify_table(doc.tables[0]) == "callout"


def test_classify_respects_manual_fallback_override():
    """If (slug, idx) is in TABLE_FALLBACK, classify_table returns 'fallback'."""
    from scripts.rebuild_from_docx.tables import TABLE_FALLBACK
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "h"
    TABLE_FALLBACK.add(("any-slug", 0))
    try:
        assert classify_table(doc.tables[0], slug="any-slug", index=0) == "fallback"
    finally:
        TABLE_FALLBACK.discard(("any-slug", 0))


def test_render_code_table_strips_known_language_hint():
    doc = Document()
    _add_consolas_table(doc, lang_first=True, body="JSON\n{\"k\": 1}")
    rendered = render_code_table(doc.tables[0])
    assert rendered == '```json\n{"k": 1}\n```'


def test_render_code_table_maps_plain_text_to_text():
    doc = Document()
    _add_consolas_table(doc, lang_first=True, body="Plain Text\ngit -v")
    assert render_code_table(doc.tables[0]) == "```text\ngit -v\n```"


def test_render_code_table_keeps_unknown_first_line_as_code():
    doc = Document()
    _add_consolas_table(doc, lang_first=True, body="本次可以执行\nls -la")
    # not a known language → no language tag, first line stays in body
    assert render_code_table(doc.tables[0]) == "```\n本次可以执行\nls -la\n```"


def test_render_data_table_to_gfm():
    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    t.rows[0].cells[0].text = "维度"
    t.rows[0].cells[1].text = "值"
    t.rows[1].cells[0].text = "速度"
    t.rows[1].cells[1].text = "快"
    t.rows[2].cells[0].text = "成本"
    t.rows[2].cells[1].text = "低"
    assert render_data_table(doc.tables[0]) == (
        "| 维度 | 值 |\n"
        "| --- | --- |\n"
        "| 速度 | 快 |\n"
        "| 成本 | 低 |"
    )


def test_render_callout_emits_paragraphs():
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].paragraphs[0].add_run("first line")
    t.rows[0].cells[0].add_paragraph("second line")
    rendered = render_callout_table(doc.tables[0])
    assert rendered == "first line\n\nsecond line"
