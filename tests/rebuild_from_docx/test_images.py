import hashlib
import struct
import zlib
from pathlib import Path

from docx import Document
from docx.shared import Inches

from scripts.rebuild_from_docx.images import extract_drawings_from_paragraph


def _make_minimal_png() -> bytes:
    """Return a valid 1x1 red pixel PNG."""
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(name: bytes, data: bytes) -> bytes:
        c = name + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)  # 1x1 8-bit RGB
    raw = b"\x00\xff\x00\x00"  # filter=0, R=255 G=0 B=0
    idat = zlib.compress(raw)
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def test_extract_image_writes_hashed_file(tmp_path: Path):
    # Build a docx with one inline image
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run()
    fixture = tmp_path / "tiny.png"
    fixture.write_bytes(_make_minimal_png())
    run.add_picture(str(fixture), width=Inches(1.0))

    out_dir = tmp_path / "out"
    out_dir.mkdir()
    refs = extract_drawings_from_paragraph(doc.paragraphs[0], doc, out_dir, slug="cognition")

    assert len(refs) == 1
    ref = refs[0]
    # ref looks like /book/cognition/<hash>.png
    assert ref.startswith("/book/cognition/")
    assert ref.endswith(".png")
    # file actually written
    written = list(out_dir.glob("*.png"))
    assert len(written) == 1
    # hash matches first 12 chars of sha256
    expected_hash = hashlib.sha256(written[0].read_bytes()).hexdigest()[:12]
    assert expected_hash in ref


def test_paragraph_with_no_drawing_returns_empty(make_doc, tmp_path: Path):
    doc = make_doc([("plain text", None)])
    out_dir = tmp_path / "out"
    out_dir.mkdir()
    assert extract_drawings_from_paragraph(doc.paragraphs[0], doc, out_dir, slug="x") == []
