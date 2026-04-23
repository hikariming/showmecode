from pathlib import Path

import pytest
from docx import Document

from scripts.rebuild_from_docx.render import render_chapter


@pytest.fixture
def anchor_md(tmp_path):
    def _write(headings: list[tuple[str, int]]) -> Path:
        md = tmp_path / "anchors.md"
        md.write_text(
            "\n\n".join(f"{'#' * lvl} {text}" for text, lvl in headings),
            encoding="utf-8",
        )
        return md
    return _write


def test_emits_anchor_heading_when_paragraph_matches(anchor_md, tmp_path):
    # H1 anchors are intentionally ignored to prevent circular drift: body content
    # rendered as H1 markdown gets re-collected as anchors on subsequent runs.
    # Only H2/H3 anchors are matched and emitted in-place.
    md = anchor_md([("第一部分：A", 1), ("Section X", 2)])
    doc = Document()
    doc.add_paragraph("intro")
    doc.add_paragraph("Section X")  # this should become "## Section X"
    doc.add_paragraph("body of section")
    out_imgs = tmp_path / "imgs"; out_imgs.mkdir()
    result = render_chapter(
        body_elements=doc.paragraphs,
        anchors_md=md,
        slug="x",
        images_dir=out_imgs,
    )
    assert "# 第一部分：A" not in result  # H1 anchors are skipped
    assert "## Section X" in result
    assert "intro" in result
    assert "body of section" in result
    # Section X must appear as a heading exactly once — never as a plain paragraph too
    assert result.count("Section X") == 1


def test_skips_front_matter_callout_table(anchor_md, tmp_path):
    md = anchor_md([("第一部分：A", 1)])
    doc = Document()
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "本章作者：明立\nAI教育工作者"
    doc.add_paragraph("real body")
    body = list(_iter_body_for_test(doc))
    out_imgs = tmp_path / "imgs"; out_imgs.mkdir()
    result = render_chapter(body, md, "x", out_imgs)
    assert "本章作者" not in result
    assert "real body" in result


def test_renders_data_table_inline(anchor_md, tmp_path):
    md = anchor_md([("第一部分：A", 1)])
    doc = Document()
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "k"
    t.rows[0].cells[1].text = "v"
    t.rows[1].cells[0].text = "1"
    t.rows[1].cells[1].text = "2"
    body = list(_iter_body_for_test(doc))
    out_imgs = tmp_path / "imgs"; out_imgs.mkdir()
    result = render_chapter(body, md, "x", out_imgs)
    assert "| k | v |" in result
    assert "| 1 | 2 |" in result


def test_h2_emitted_in_body_position_not_at_top(anchor_md, tmp_path):
    """H2 anchors must be emitted where they match in the body, not stacked at the top.
    H1 anchors are ignored entirely to prevent circular drift on re-runs."""
    md = anchor_md([("Title", 1), ("First Section", 2), ("Second Section", 2)])
    doc = Document()
    doc.add_paragraph("intro before any section")
    doc.add_paragraph("First Section")
    doc.add_paragraph("body of first")
    doc.add_paragraph("Second Section")
    doc.add_paragraph("body of second")
    body = list(_iter_body_for_test(doc))
    out_imgs = tmp_path / "imgs"; out_imgs.mkdir()
    result = render_chapter(body, md, "x", out_imgs)
    # H1 "Title" anchor is skipped; output starts with body content.
    # Order: intro, then first heading in-place, body, second heading in-place, body.
    # The two H2s must NOT be adjacent (body text between them proves in-place emission).
    lines = [l for l in result.splitlines() if l.strip()]
    assert "# Title" not in result  # H1 anchors are skipped
    assert lines[0] == "intro before any section"
    assert lines[1] == "## First Section"
    assert lines[2] == "body of first"
    assert lines[3] == "## Second Section"
    assert lines[4] == "body of second"


def test_anchor_match_tolerates_whitespace_drift(anchor_md, tmp_path):
    """Whitespace differences between .md anchor text and docx paragraph text
    must not silently drop the heading. The anchor's wording wins (it's the
    peer-reviewed source of truth)."""
    md = anchor_md([("AI是乙方", 3)])
    doc = Document()
    doc.add_paragraph("AI 是乙方")  # extra space
    body = list(_iter_body_for_test(doc))
    out_imgs = tmp_path / "imgs"; out_imgs.mkdir()
    result = render_chapter(body, md, "x", out_imgs)
    assert "### AI是乙方" in result
    assert "AI 是乙方" not in result  # drift wording must NOT appear


def test_anchor_match_strips_arabic_numbering(anchor_md, tmp_path):
    """docx body says '1. 言出法随', .md anchor says '言出法随' — should match."""
    md = anchor_md([("言出法随", 2)])
    doc = Document()
    doc.add_paragraph("1. 言出法随")
    body = list(_iter_body_for_test(doc))
    out_imgs = tmp_path / "imgs"; out_imgs.mkdir()
    result = render_chapter(body, md, "x", out_imgs)
    assert "## 言出法随" in result
    # docx wording must NOT appear
    assert "1. 言出法随" not in result


def test_anchor_match_strips_chinese_numbering(anchor_md, tmp_path):
    """docx body says '2、总结', .md anchor says '总结' — should match."""
    md = anchor_md([("总结", 2)])
    doc = Document()
    doc.add_paragraph("2、总结")
    body = list(_iter_body_for_test(doc))
    out_imgs = tmp_path / "imgs"; out_imgs.mkdir()
    result = render_chapter(body, md, "x", out_imgs)
    assert "## 总结" in result
    assert "2、总结" not in result


def _iter_body_for_test(doc):
    """Mirror slicer's body-iteration so tests can hand a chapter slice to render_chapter."""
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    from docx.table import Table
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)
