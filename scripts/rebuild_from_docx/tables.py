"""Classify and render docx tables.

Three classes seen in 《赤脚程序员实战手册》.docx (244 tables total):

- code: 1×1 cell whose runs are all Consolas. First text line is a language hint.
- data: ≥2 rows × ≥2 cols, no Consolas runs.
- callout: anything else (1×1 plain prose, 1×N author/intro boxes).
"""
from docx.table import Table

from .runs import paragraph_to_markdown

CODE_FONTS = {"Consolas", "Menlo", "Courier", "Courier New", "Monaco"}

# Manual escape hatch: tables that cannot render cleanly as GFM (merged cells,
# images-in-cells, very wide content). Keyed by (slug, table_index_within_chapter).
# Empty by default; populated only if T8 manual gate uncovers a broken table.
TABLE_FALLBACK: set[tuple[str, int]] = set()

LANG_MAP = {
    "plain text": "text",
    "shell": "bash",
    "ps": "powershell",
    "powershell": "powershell",
    "bash": "bash",
    "sh": "bash",
    "json": "json",
    "yaml": "yaml",
    "yml": "yaml",
    "javascript": "javascript",
    "js": "javascript",
    "typescript": "typescript",
    "ts": "typescript",
    "html": "html",
    "css": "css",
    "sql": "sql",
    "python": "python",
    "py": "python",
    "markdown": "markdown",
    "md": "markdown",
    "swift": "swift",
}


def _all_consolas(table: Table) -> bool:
    """True if every non-whitespace run in every cell is in a code font."""
    saw_any = False
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if not r.text.strip():
                        continue
                    saw_any = True
                    if (r.font.name or "") not in CODE_FONTS:
                        return False
    return saw_any


def classify_table(table: Table, slug: str | None = None, index: int | None = None) -> str:
    if slug is not None and index is not None and (slug, index) in TABLE_FALLBACK:
        return "fallback"
    rows = len(table.rows)
    cols = len(table.columns) if rows else 0
    if rows == 1 and cols == 1 and _all_consolas(table):
        return "code"
    if rows >= 2 and cols >= 2 and not _all_consolas(table):
        return "data"
    return "callout"


def render_fallback_placeholder(slug: str, index: int) -> str:
    """Emit an HTML comment for the author to replace with a hand-pasted screenshot."""
    return f"<!-- TABLE FALLBACK: {slug} table #{index} — paste a screenshot here -->"


def render_code_table(table: Table) -> str:
    cell = table.rows[0].cells[0]
    raw = "\n".join(p.text for p in cell.paragraphs).rstrip()
    lines = raw.split("\n")
    lang = ""
    body_lines = lines
    if lines:
        first = lines[0].strip().lower()
        if first in LANG_MAP:
            lang = LANG_MAP[first]
            body_lines = lines[1:]
    body = "\n".join(body_lines).rstrip()
    return f"```{lang}\n{body}\n```"


def render_data_table(table: Table) -> str:
    rows_text = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
    if not rows_text:
        return ""
    header = "| " + " | ".join(rows_text[0]) + " |"
    sep = "| " + " | ".join(["---"] * len(rows_text[0])) + " |"
    body = ["| " + " | ".join(row) + " |" for row in rows_text[1:]]
    return "\n".join([header, sep, *body])


def render_callout_table(table: Table) -> str:
    """Unwrap a 1×N callout: emit each cell's paragraphs as plain markdown, blank-line separated."""
    parts: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                text = paragraph_to_markdown(p)
                if text.strip():
                    parts.append(text)
    return "\n\n".join(parts)
